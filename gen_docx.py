
# -*- coding: utf-8 -*-
import os, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

for level in range(1, 4):
    hs = doc.styles["Heading %d" % level]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sizes = {1: 16, 2: 13, 3: 11}
    hs.font.size = Pt(sizes[level])

sec = doc.sections[0]
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)

def ap(text, bold=False, sz=10, al=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    r.font.size = Pt(sz)
    if al is not None: p.alignment = al
    p.paragraph_format.space_after = Pt(6)
    return p

def alink(text, url, sz=9):
    p = doc.add_paragraph()
    r1 = p.add_run(text + ": ")
    r1.font.size = Pt(sz)
    r1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r2 = p.add_run(url)
    r2.font.size = Pt(sz)
    r2.font.color.rgb = RGBColor(0x33, 0x66, 0xcc)
    r2.underline = True
    p.paragraph_format.space_after = Pt(2)

def bullet(text, sz=9):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(sz)
    p.paragraph_format.space_after = Pt(2)

with open(r"D:\everyday_recommand\weekly_data.json", "r", encoding="utf-8") as f:
    data = json.load(f)

# Title
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run(data["title"])
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
tp.paragraph_format.space_after = Pt(4)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run(data["subtitle"])
sr.font.size = Pt(11)
sr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
sp.paragraph_format.space_after = Pt(6)

gp = doc.add_paragraph()
gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
gr = gp.add_run(data["gen_date"])
gr.font.size = Pt(8)
gr.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
gp.paragraph_format.space_after = Pt(16)

# Section 1
doc.add_heading(data["sections"][0]["heading"], level=2)
ap(data["sections"][0]["body"], sz=10)

# Section 2 - Papers
doc.add_heading(data["sections"][1]["heading"], level=2)
for paper in data["sections"][1]["papers"]:
    doc.add_heading(paper["title_en"], level=3)
    ap(paper["meta"], bold=True, sz=9)
    if paper.get("doi"):
        ap("DOI: " + paper["doi"], sz=9)
    ap(paper["desc"], sz=10)

# Recommended
doc.add_heading(data["sections"][1].get("rec_heading", "Recommended Reading"), level=3)
for rec in data["sections"][1].get("recommended", []):
    bullet(rec, sz=9)

# Section 3 - Enterprise
doc.add_heading(data["sections"][2]["heading"], level=2)
for ent in data["sections"][2]["items"]:
    doc.add_heading(ent["title"], level=3)
    ap("Date: " + ent["date"], bold=True, sz=9)
    ap(ent["desc"], sz=10)
    alink("Source", ent["url"])

# Section 4 - Policy
doc.add_heading(data["sections"][3]["heading"], level=2)
for pol in data["sections"][3]["items"]:
    doc.add_heading(pol["title"], level=3)
    ap("Date: " + pol["date"], bold=True, sz=9)
    ap(pol["desc"], sz=10)
    alink("Source", pol["url"])

# Section 5 - Conferences
doc.add_heading(data["sections"][4]["heading"], level=2)
for conf in data["sections"][4]["items"]:
    p = doc.add_paragraph()
    r = p.add_run("\u2022 " + conf["name"])
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(" \u2014 " + conf["date"])
    p.paragraph_format.space_after = Pt(2)
    alink("  URL", conf["url"], sz=9)

# Section 6 - Outlook
doc.add_heading(data["sections"][5]["heading"], level=2)
ap(data["sections"][5]["body"], sz=10)
for o in data["sections"][5]["items"]:
    bullet(o, sz=10)

# Footer
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(data["footer"])
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
fr.italic = True

out_dir = r"D:\everyday_recommand\reports\2026\07"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "2026-07-24-??.docx")
doc.save(out_path)
print("DOCX saved: " + out_path)
print("Size: " + str(os.path.getsize(out_path)) + " bytes")
