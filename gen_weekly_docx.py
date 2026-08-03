# -*- coding: utf-8 -*-
 """鐢熸垚閲忓瓙鍏夊绮惧瘑娴嬮噺鍛ㄦ姤 DOCX"""
 import os
 from docx import Document
 from docx.shared import Pt, Inches, Cm, RGBColor
 from docx.enum.text import WD_ALIGN_PARAGRAPH
 from docx.enum.style import WD_STYLE_TYPE
 from docx.oxml.ns import qn
 
 doc = Document()
 
 # -- style setup --
 style = doc.styles['Normal']
 font = style.font
 font.name = 'Calibri'
 font.size = Pt(10.5)
 style.element.rPr.rFonts.set(qn('w:eastAsia'), '寰蒋闆呴粦')
 
 for level in range(1, 4):
     hstyle = doc.styles[f'Heading {level}']
     hfont = hstyle.font
     hfont.name = 'Calibri'
     hfont.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
     hstyle.element.rPr.rFonts.set(qn('w:eastAsia'), '寰蒋闆呴粦')
     if level == 1:
         hfont.size = Pt(16)
         hfont.bold = True
     elif level == 2:
         hfont.size = Pt(13)
         hfont.bold = True
     else:
         hfont.size = Pt(11)
         hfont.bold = True
 
 section = doc.sections[0]
 section.top_margin = Cm(2.0)
 section.bottom_margin = Cm(2.0)
 section.left_margin = Cm(2.2)
 section.right_margin = Cm(2.2)
 
 def add_para(text, bold=False, size=None, align=None, space_after=6):
     p = doc.add_paragraph()
     run = p.add_run(text)
     if bold: run.bold = True
     if size: run.font.size = Pt(size)
     if align is not None: p.alignment = align
     p.paragraph_format.space_after = Pt(space_after)
     p.paragraph_format.space_before = Pt(0)
     return p
 
 def add_link_para(text, url, size=9):
     p = doc.add_paragraph()
     run = p.add_run(f'{text}: ')
     run.font.size = Pt(size)
     run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
     r2 = p.add_run(url)
     r2.font.size = Pt(size)
     r2.font.color.rgb = RGBColor(0x33, 0x66, 0xcc)
     r2.underline = True
     p.paragraph_format.space_after = Pt(2)
     return p
 
 # ============ TITLE ============
 title_p = doc.add_paragraph()
 title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
 title_run = title_p.add_run('閲忓瓙鍏夊绮惧瘑娴嬮噺 路 宸ヤ綔杩涘睍鍛ㄦ姤')
 title_run.bold = True
 title_run.font.size = Pt(20)
 title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
 title_p.paragraph_format.space_after = Pt(4)
 
 sub_p = doc.add_paragraph()
 sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
 sub_run = sub_p.add_run('绗?30 鍛?路 2026-07-20 鈥?2026-07-24')
 sub_run.font.size = Pt(11)
 sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
 sub_p.paragraph_format.space_after = Pt(6)
 
 gen_p = doc.add_paragraph()
 gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
 gen_run = gen_p.add_run('鐢熸垚鏃ユ湡锛?026-07-24 | 鑷姩鍖栫敓鎴?)
 gen_run.font.size = Pt(8)
 gen_run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
 gen_p.paragraph_format.space_after = Pt(16)
 
 # ============ 涓€銆佹湰鍛ㄧ儹鐐?============
 doc.add_heading('涓€銆佹湰鍛ㄧ儹鐐?, level=2)
 add_para(
     '鏈懆閲忓瓙鍏夊绮惧瘑娴嬮噺棰嗗煙鍦ㄧ鐮斻€佷紒涓氥€佹斂绛栦笁涓淮搴﹀潎鏈夐噸瑕佸姩鎬併€傜鐮旀柟闈紝閲忓瓙闄€铻轰华鍙岃矾寰勬帹杩?
     '锛堥噺瀛愬寮哄帇缂╁厜鏉熶笌瓒呭ぇ琚姩璋愭尟鑵旓級鍜?SU(1,1) 骞叉秹浠娆＄粨鍚?Rydberg 鍘熷瓙瓒呰秺 SQL 鏄袱澶т寒鐐癸紱'
     '姝ゅ锛屽樊鍒嗗師瀛愬共娑変华鐨勫師鍨嬮獙璇侊紙Nature锛夊拰鍋忔尟鍏煎閲忓瓙瀛樺偍锛圕omm Phys锛変负寮曞姏娉㈡帰娴嬪拰閲忓瓙缃戠粶鎻愪緵浜?
     '鍏抽敭鎶€鏈獙璇併€備紒涓氭柟闈紝Q-CTRL 鍦ㄨ寖鍫＄綏鑸睍鎺ㄥ嚭鑸┖閫傝埅绾ч噺瀛愬鑸郴缁燂紝QuantumDiamonds 鑾峰緱 9100 涓囨鍏?
     '铻嶈祫锛孖nfleqtion 涓庤嫳鍥界殗瀹舵捣鍐涙紨绀烘按涓嬮噺瀛愬厜閽熴€傛斂绛栨柟闈紝缇庡浗鎬荤粺绛剧讲閲忓瓙鍒涙柊琛屾斂浠わ紙EO 14413锛夛紝'
     'NSF 鍚姩 Project Triad锛屽浗闃查儴 DIU 鎷ㄦ 2 浜跨編鍏冨惎鍔?Farseer 閲忓瓙浼犳劅璁″垝锛岃繖浜涙斂绛栧姩浣滃皢鏄捐憲鍔犻€?
     '閲忓瓙浼犳劅浠庡疄楠屽鍒伴儴缃插簲鐢ㄧ殑杞寲銆?,
     size=10
 )
 
 # ============ 浜屻€佺鐮旇繘灞?============
 doc.add_heading('浜屻€佺鐮旇繘灞?, level=2)
 
 # Paper 1
 doc.add_heading('1. Enhanced phase estimation with coherently boosted two-mode squeezed beams and its application to optical gyroscopes', level=3)
 add_para('arXiv: 2607.05732 | 2026-07-07 | Xiao-Qi Xiao, Elisha S. Matekole, Jiankang Zhao et al.', bold=True, size=9)
 add_para(
     '鍦ㄥ厜瀛﹂檧铻轰华涓彁鍑虹浉骞插寮哄弻妯″帇缂╁厜鏉熺殑璁￠噺鏂规锛岀悊璁鸿瘉鏄庡彲瀹炵幇瓒呰秺缁忓吀鏋侀檺鐨勭浉浣嶄及璁＄簿搴︺€?
     '璇ュ伐浣滅洿鎺ュ搴旈噺瀛愬厜绾ら檧铻轰华鏂瑰悜锛屾彁渚涗簡鍘嬬缉鍏夊湪鏃嬭浆浼犳劅涓殑鏂版柟娉曘€?,
     size=10
 )
 
 # Paper 2
 doc.add_heading('2. Standard-quantum-limit-surpassing vector polarimetry using Rydberg atoms in an SU(1,1) interferometer', level=3)
 add_para('arXiv: 2606.27870 | 2026-06-26 | Weiqiang Guan, Yuetao Chen, Jun Zhou et al.', bold=True, size=9)
 add_para(
     '灏?Rydberg 鍘熷瓙浼犳劅涓?SU(1,1) 闈炵嚎鎬у共娑変华缁撳悎锛屽疄楠屽睍绀轰簡瓒呰秺鏍囧噯閲忓瓙鏋侀檺鐨勭煝閲忓亸鎸祴閲忋€?
     '杩欐槸 SU(1,1) 骞叉秹浠娆＄粨鍚?Rydberg 鍘熷瓙骞惰秴瓒?SQL锛屽叿鏈夐噷绋嬬鎰忎箟銆?,
     size=10
 )
 
 # Paper 3
 doc.add_heading('3. Robust phase sensitivity in Mach-Zehnder interferometer using photon added and subtracted squeezed coherent state', level=3)
 add_para('arXiv: 2607.02984 | 2026-07-03 | Shivani Singh, Priya Malpani, Anirban Pathak et al.', bold=True, size=9)
 add_para(
     '鐞嗚鐮旂┒鍏夊瓙澧炲噺鎿嶄綔浣滅敤浜庡帇缂╃浉骞叉€佸悗鍦?MZI 涓殑鐩镐綅鐏垫晱搴︼紝灞曠幇鍑烘瘮浼犵粺鍘嬬缉鎬佹洿浼樼殑鎬ц兘鍜?
     '鎹熻€楅瞾妫掓€с€備负閲忓瓙骞叉秹浠殑闈炵粡鍏告€侀€夋嫨鎻愪緵浜嗘柊鍊欓€夈€?,
     size=10
 )
 
 # Paper 4
 doc.add_heading('4. Restoring Velocity Immunity via Dynamic Mirror Compensation in a Large-Area Dual-Atom-Interferometer Gyroscope', level=3)
 add_para('arXiv: 2605.29929 | 2026-05-28 | Jie Gu, Yin-fei Mao, Zhan-Wei Yao et al.', bold=True, size=9)
 add_para(
     '鎻愬嚭鍔ㄦ€佸弽灏勯暅琛ュ伩鏂规鎭㈠澶ч潰绉弻鍘熷瓙骞叉秹浠檧铻轰华鐨勯€熷害鍏嶇柅鎬э紝瀹為獙楠岃瘉浜嗙簿搴︽樉钁楁彁鍗囥€?
     '瑙ｅ喅浜嗗師瀛愬共娑夐檧铻轰华鐨勫叧閿伐绋嬫寫鎴樷€斺€旈€熷害鍏嶇柅涓уけ闂銆?,
     size=10
 )
 
 # Paper 5
 doc.add_heading('5. Ultralow shot noise limited giant passive resonant gyroscope for Earth rotation measurement', level=3)
 add_para('arXiv: 2606.06822 | 2026-06-05 | Yuhong Zhong, Yangsheng Cai, Zhanhao Liu et al.', bold=True, size=9)
 add_para(
     '鎶ラ亾瓒呭ぇ鐜舰鑵旇鍔ㄨ皭鎸檧铻轰华锛屾暎绮掑櫔澹版瀬闄愯揪鍒版瀬浣庢按骞筹紝鍙洿鎺ユ祴閲忓湴鐞冭嚜杞紝'
     '涓洪噺瀛愬寮烘柟妗堟彁渚涗簡瀵规瘮鍩哄噯銆傚叾鍣０鏋侀檺瀵硅瘎浼伴噺瀛愰檧铻轰华鐨勫疄鐢ㄤ紭鍔胯嚦鍏抽噸瑕併€?,
     size=10
 )
 
 # Paper 6 (from daily 7/24)
 doc.add_heading('6. [娣卞害鍒嗘瀽] A prototype differential atom interferometer for fundamental physics', level=3)
 add_para('Nature Vol 654, 2026-06-17 | C. F. A. Baynham, R. Hobson, O. Buchm眉ller et al. (AION Collaboration)', bold=True, size=9)
 add_para('DOI: 10.1038/s41586-026-10617-1', size=9)
 add_para(
     'AION 鍚堜綔缁勫湪瀹為獙瀹ゆ惌寤轰簡鍩轰簬璐圭背瀛?鈦糕伔Sr 鍗曞厜瀛愰挓璺冭縼鐨勫樊鍒嗗師瀛愬共娑変华鍘熷瀷銆傛牳蹇冮獙璇侊細'
     '锛?锛夊樊鍒嗘瀯鍨嬪湪娉ㄥ叆鏁板姬搴﹀悎鎴愭縺鍏夌浉浣嶅櫔澹颁笅浠嶇淮鎸侀噺瀛愭瀬闄愮伒鏁忓害锛?
     '锛?锛夊湪瀹屽叏鐩镐綅闅忔満鍖栨潯浠朵笅鎭㈠鐩稿共鎸崱淇″彿銆傞娆＄敤 鈦糕伔Sr 楠岃瘉宸垎鍣０鎶戝埗锛?
     '涓哄叕閲岀骇鍩虹嚎涔冭嚦绌洪棿 AEDGE 寮曞姏娉㈡帰娴嬫彁渚涗簡鍏抽敭瀹為獙渚濇嵁銆?,
     size=10
 )
 
 # Paper 7 (from daily 7/24)
 doc.add_heading('7. [娣卞害鍒嗘瀽] Full-polarization-compatible optical storage for arbitrary polarized modes', level=3)
 add_para('Communications Physics, 2026-07-17 | Xin Yang, Jinwen Wang, Zehong Chang et al. (瑗垮畨浜ら€氬ぇ瀛?', bold=True, size=9)
 add_para('DOI: 10.1038/s42005-026-02769-3', size=9)
 add_para(
     '鍦ㄥ喎 鈦糕伔Rb 鍘熷瓙绯荤患涓紝閲囩敤璺緞涓嶅垎绂绘灦鏋勫疄鐜颁簡瀵逛换鎰忓亸鎸ā寮忕殑楂樻晥銆佸亸鎸笉鏁忔劅鍏夊瓨鍌ㄣ€?
     '瀛樺偍鏁堢巼瓒呰繃 50%锛屽瓨鍌ㄤ繚鐪熷害瓒呰繃 95%锛堣繙瓒呯粡鍏告瀬闄愶級銆傛鏂规娑堥櫎浜嗕紶缁?EIT 閲忓瓙瀛樺偍涓?
     '鍋忔尟-璺緞杞崲鐨勭摱棰堬紝瀵归噺瀛愪腑缁у櫒鎺ュ彛璁捐鏈夐噸瑕佹剰涔夈€?,
     size=10
 )
 
 # Paper 8 (from daily 7/23)
 doc.add_heading('8. [娣卞害鍒嗘瀽] Magnetometry with a space-based differential atom interferometer', level=3)
 add_para('Nature Communications 17, 2026-07-11 | Matthias Meister, Gabriel M眉ller et al. (DLR/JPL/UC Berkeley)', bold=True, size=9)
 add_para('DOI: 10.1038/s41467-026-75230-2', size=9)
 add_para(
     '鍦?NASA Cold Atom Lab锛圛SS锛変腑锛屼娇鐢?鈦糕伔Rb BEC 閫氳繃鎬佸伐绋嬪皢 Bragg 鑴夊啿鏁堢巼鎻愬崌鑷虫敮鎾?'
     '2T = 40.3 ms 鐨勫共娑夋椂闂达紙姝ゅ墠浠?4 ms锛夈€傞娆″疄鐜扮┖闂村樊鍒?BEC 纾佸姏璁★細MZ 鍨嬫祴鍔犻€熷害姊害锛?
     'butterfly 鍨嬫祴鍔犻€熷害鏇茬巼銆傜鍔涜涓嶇‘瀹氬害杈?~5 nT/mm虏锛屾瘮缁忓吀 COM 杩愬姩娉曠簿搴︽彁鍗囩害 57 鍊嶃€?
     'BEC 鑶ㄨ儉娓╁害浣庤嚦 0.66 nK銆?,
     size=10
 )
 
 # Recommended reading
 doc.add_heading('鎺ㄨ崘闃呰', level=3)
 recs = [
     'Quantum Sensors for Chemistry and Materials Science [arXiv:2607.07848, 2026-07-08] 鈥?缁艰堪閲忓瓙浼犳劅鍣ㄥ湪鍖栧鍜屾潗鏂欑瀛︿腑鐨勫簲鐢紝瑕嗙洊 NV 鑹插績銆佸師瀛愮鍔涜绛夊钩鍙?,
     'Local, mm-scale 鹿H magnetic resonance imaging using atomic vapors [arXiv:2606.15946, 2026-06-14] 鈥?鍒╃敤鍏夋车纾佸姏璁″疄鐜?渭L 绾ф牱鍝佺殑浣庡満 MRI',
     'A Signal Analysis Framework for Unshielded Room-Temperature Magnetocardiography [arXiv:2606.29285, 2026-06-28] 鈥?瀹ゆ俯闈炲睆钄藉績纾佸浘淇″彿鍒嗘瀽妗嗘灦锛孫PM 鐢熺墿纾佹祴閲忓簲鐢?,
     'Optical nanofiber testbeds for on-chip quantum inertial sensing [AVS Quantum Science, 2026-07-16] 鈥?鍊忛€濆満鍏夌撼绾ら獙璇佺墖涓婇噺瀛愬姞閫熷害璁″拰闄€铻轰华',
     'Transit Noise in Spin Squeezing Experiments with Coated Rubidium Vapor Cell [Photonics, 2026-05-06] 鈥?鐑師瀛愯嚜鏃嬪帇缂╀腑 transit noise 瀵规€ц兘鐨勯檺鍒剁爺绌?,
     'Wafer-scale micro-knife sealed vacuum cells for quantum devices [Quantum Sci. Technol., 2026-07-16] 鈥?NIST 鏅跺渾绾у井鍒€鍙ｉ敭鍚堢湡绌鸿厰宸ヨ壓锛岄€傜敤浜?OPM/鍘熷瓙閽熻姱鐗囧寲灏佽',
     'Noise Mitigation in Quantum-Enhanced Fiber Optic Gyroscopes [Quantum Reports, 2026-05-01] 鈥?閲忓瓙澧炲己鍏夌氦闄€铻轰华涓殑鍣０鎶戝埗',
 ]
 for r in recs:
     p = doc.add_paragraph()
     p.style = doc.styles['List Bullet']
     run = p.add_run(r)
     run.font.size = Pt(9)
     p.paragraph_format.space_after = Pt(2)
 
 # ============ 涓夈€佷紒涓氬姩鎬?============
 doc.add_heading('涓夈€佷紒涓氬姩鎬?, level=2)
 
 enterprises = [
     ('1. Q-CTRL 鍦ㄨ寖鍫＄綏鑸睍鎺ㄥ嚭鑸┖閫傝埅绾ч噺瀛?GPS 澶囦唤绯荤粺',
      '2026-07-15鈥?9',
      'Q-CTRL 鍦ㄨ寖鍫＄綏鍥介檯鑸睍灞曠ず浜嗗叾 Ironstone Opal 閲忓瓙瀵艰埅瑙ｅ喅鏂规锛屽凡鑾峰緱鑸┖閫傝埅璧勮川銆?
      '绌哄鍜屾礇椹潎涓哄悎浣滀紮浼达紝鐩爣甯傚満娑电洊姘戠敤鑸┖銆佸浗闃插拰鑷姩椹鹃┒銆侴PS 鎷掓鐜涓嬬殑閲忓瓙瀵艰埅'
      '鍟嗕笟鍖栬繄鍑哄叧閿竴姝ャ€?,
      'https://www.ainonline.com/aviation-news/aerospace/2026-07-18/q-ctrl-eyes-multiple-markets-quantum-navigation-system'),
     ('2. QuantumDiamonds 鑾峰緱 9100 涓囨鍏冭瀺璧勶紝鎺ㄨ繘閲戝垰鐭冲崐瀵间綋妫€娴?,
      '2026-07-16',
      '鎱曞凹榛戝垱涓氬叕鍙?QuantumDiamonds 瀹屾垚澶ч铻嶈祫锛屼笓娉ㄤ簬鍩轰簬 NV 鑹插績鐨勯噺瀛愮簿瀵嗘祴閲忔妧鏈紝'
      '鐢ㄤ簬鍗婂浣撴棤鎹熸娴嬨€傞噾鍒氱煶閲忓瓙浼犳劅鍣ㄥ湪宸ヤ笟搴旂敤涓殑鍟嗕笟鍖栨鍦ㄥ姞閫熴€?,
      'https://techbriefe.com/tech/startups/quantumdiamonds-secures-91-million-for-semiconductor-innovation'),
     ('3. Infleqtion 涓庤嫳鍥界殗瀹舵捣鍐涙紨绀轰笘鐣岄涓按涓嬮噺瀛愬厜閽?,
      '2026-07-17',
      'Infleqtion 鐨?Tiqker 閲忓瓙鍏夐挓鎴愬姛鍦ㄨ嫳鍥界殗瀹舵捣鍐?XCal 鑷富娼滆墖涓婃紨绀猴紝瀹炵幇 GPS 鎷掓鐜涓嬬殑'
      '姘翠笅瀵艰埅銆傛澶?Infleqtion 鍚屾椂鑾峰緱 DOE 300 涓囩編鍏冭祫鍔╃敤浜庨噺瀛愪紶鎰熶笌璁＄畻銆?,
      'https://cyprusshippingnews.com/2026/07/17/infleqtion-and-royal-navy-demonstrate-worlds-first-quantum-optical-clock-on-underwater-autonomous-submarine-to-advance-gps-free-navigation/'),
     ('4. Honeywell 鐗靛ご涓?ESA 寮€鍙戠揣鍑戝瀷閲忓瓙纾佸姏璁?,
      '2026-07-15',
      'Honeywell Aerospace 涓?Quantum Brilliance銆佹尝鍏伴泤鐩栨矁澶у鍚堜綔锛屼负娆ф床绌洪棿灞€寮€鍙戙€佹祴璇曞苟浜や粯'
      '绱у噾鍨嬮噺瀛愮鍔涜銆?,
      'https://www.militaryaerospace.com/sensors/article/55390887/honeywell-led-group-to-develop-compact-quantum-magnetometer-for-esa'),
     ('5. Quantum X Labs 鍙戝竷棣栦釜鍏ㄥ厜瀛﹀崐鐞冭皭鎸檧铻轰华',
      '2026-07-15',
      '閲忓瓙鎶€鏈叕鍙?Quantum X Labs 鍙戝竷浜嗛涓叏鍏夊鍗婄悆璋愭尟闄€铻轰华锛岄噰鐢ㄥ叏鍏夊璇诲嚭鏂规銆?,
      'https://www.gophotonics.com/news/details/9056-quantum-x-labs-unveils-first-fully-all-optical-hemispherical-resonator-gyroscope'),
     ('6. Photon Queue 鑾峰緱 400 涓囩編鍏冪瀛愯疆铻嶈祫',
      '2026-07-21',
      '閲忓瓙瀛樺偍鍒涗笟鍏徃 Photon Queue 瀹屾垚瓒呴璁よ喘鐨?400 涓囩編鍏冪瀛愯疆铻嶈祫锛岃嚧鍔涗簬瑙ｅ喅閲忓瓙璁＄畻涓殑'
      '鍏夊瓙瀛樺偍鐡堕銆?,
      'https://techstartups.com/2026/07/21/photon-queue-raises-4m-seed-round-to-advance-quantum-memory-for-next-generation-quantum-computing/'),
     ('7. 浣庡姛鑰?Ramsey 閲戝垰鐭抽噺瀛愮鍔涜鐢ㄤ簬鐢熺墿纾佸満娴嬮噺',
      '2026-07-22',
      '涓滀含鐞嗙澶у鐮斿彂浜嗗熀浜?Ramsey 搴忓垪鐨勪綆鍔熻€楅噾鍒氱煶閲忓瓙纾佸姏璁★紝鍙畨鍏ㄦ娴嬩汉浣撶敓鐗╃鍦轰俊鍙枫€?,
      'https://phys.org/news/2026-07-body-magnetic-fields-power-ramsey.html'),
 ]
 
 for title, date, desc, url in enterprises:
     doc.add_heading(title, level=3)
     add_para(f'鏃ユ湡锛歿date}', bold=True, size=9)
     add_para(desc, size=10)
     add_link_para('鏉ユ簮', url)
 
 # ============ 鍥涖€佹斂绛栨洿鏂?============
 doc.add_heading('鍥涖€佹斂绛栨洿鏂?, level=2)
 
 policies = [
     ('1. 缇庡浗鎬荤粺绛剧讲閲忓瓙鍒涙柊琛屾斂浠わ紙EO 14413锛?,
      '2026-06-22',
      '鐗规湕鏅斂搴滅缃茶鏀夸护 "Ushering in the Next Frontier of Quantum Innovation"锛岃姹傝仈閭﹀悇鏈烘瀯鍦?'
      '180 澶╁唴鍒跺畾缁煎悎閲忓瓙鎶€鏈垬鐣ワ紝灏嗙爺绌堕鍏堣浆鍖栦负閮ㄧ讲鍜屽晢涓氳妯°€傝浠ょ洿鎺ユ帹鍔ㄤ簡鍥介槻閮ㄩ噺瀛愪紶鎰?
      '椤圭洰鐨勫惎鍔ㄣ€?,
      'https://www.whitehouse.gov/presidential-actions/2026/06/ushering-in-the-next-frontier-of-quantum-innovation/'),
     ('2. 缇庡浗鍥介槻閮?DIU 鍚姩 2 浜跨編鍏?"Farseer" 閲忓瓙浼犳劅璁″垝',
      '2026-07-02',
      '鍥介槻鍒涙柊鍗曞厓鍚姩澶氶樁娈甸噺瀛愪紶鎰熷拰鏃堕璁″垝锛屾€婚绠?2 浜跨編鍏冿紝鐩爣鏄儴缃插啗姘戠敤閲忓瓙浼犳劅鍜屽畾鏃剁‖浠躲€?
      '杩欐槸鐩墠鍏紑鎶ラ亾涓妯℃渶澶х殑閲忓瓙浼犳劅涓撻」璧勫姪銆?,
      'https://quantumcomputingreport.com/department-of-wars-diu-launches-200-million-farseer-initiative-to-field-dual-use-quantum-sensing-and-timing-hardware/'),
     ('3. NSF 鍚姩 Project Triad',
      '2026-07-07',
      '缇庡浗鍥藉绉戝鍩洪噾浼氬惎鍔ㄩ涓泦鎴愰噺瀛愪紶鎰熴€侀噺瀛愮綉缁滃拰閲忓瓙璁＄畻鐨勮鍒?Project Triad锛屾帹鍔ㄩ噺瀛愭妧鏈悜'
      '瀹為檯搴旂敤杞寲銆?,
      'https://www.nsf.gov/news/nsf-launches-project-triad-advance-quantum-technology-real'),
     ('4. 缇庡浗鑳芥簮閮ㄥ惎鍔?Quantum Genesis 璁″垝',
      '2026-06-22鈥?3',
      '鑳芥簮閮ㄧ户琛屾斂浠ゅ悗鍚姩 Quantum Genesis 璁″垝锛岃嚧鍔涗簬鍒涘缓涓栫晫棣栦釜绉戝鐩稿叧鐨勫閿欓噺瀛愯绠楁満銆?
      '鍚屾椂 DOE 璧勫姪 Infleqtion 绛変紒涓氱殑閲忓瓙浼犳劅椤圭洰锛?00 涓囩編鍏冿級銆?,
      'https://www.energy.gov/science/articles/energy-department-announces-initiative-create-and-deploy-worlds-first'),
     ('5. DARPA RoQS 璁″垝锛歂V 閲戝垰鐭抽噺瀛愪紶鎰熷櫒鍐涗簨搴旂敤',
      '2026-06',
      'DARPA 鐨?Robust Quantum Sensors 璁″垝宸茬‘瀹氬畬鏁存壙鍖呭晢鍩哄湴锛岄噸鐐规帹杩?NV 閲戝垰鐭抽噺瀛愮鍔涜鍦?
      '鍐涗簨纾佸紓甯告娴嬪拰鍙嶆綔浣滄垬涓殑搴旂敤銆?,
      'https://lyceumintelligence.com/'),
     ('6. 娆ф床 QOMPASS 鑱旂洘鍛煎悂寤虹珛涓撳睘娆ф床閲忓瓙璁″垝',
      '2026-07-20鈥?2',
      'QOMPASS 鑱旂洘鍙戝竷绔嬪満鏂囦欢锛屾暒淇冩濮斾細寤虹珛涓撳睘鐨勬娲查噺瀛愯鍒掞紝骞朵笌 Horizon Europe 鐨勬娲?
      '鍚堜綔浼欎即鍏崇郴瀹℃煡鍚屾銆傞噺瀛愭棗鑸拌鍒掓鍦ㄤ妇鍔炲鍦鸿矾绾垮浘鐮旇浼氥€?,
      'https://quantumzeitgeist.com/european-quantum-initiative-qompass-consortium/'),
     ('7. 涓浗鍔犲ぇ閲忓瓙鎶€鏈祫鏈姩鍛樺拰鐩戠鎺ㄨ繘',
      '2026-06-30',
      '涓浗姝ｅ湪蹇€熸暣鍚堝浗鍐呴噺瀛愭妧鏈敓鎬佺郴缁燂紝鍔ㄥ憳鏂拌祫鏈苟鍒跺畾鐩稿叧娉曡锛屾帹鍔ㄩ噺瀛愯绠楀熀纭€璁炬柦寤鸿銆?
      '杩欎篃灏嗛棿鎺ュ奖鍝嶉噺瀛愮簿瀵嗘祴閲忕殑鍥藉唴鐮斿彂鏍煎眬銆?,
      'https://www.chinatechnews.com/2026/06/30/124453-china-mobilizes-new-capital-and-regulations-to-supercharge-quantum-computing-infrastructure'),
 ]
 
 for title, date, desc, url in policies:
     doc.add_heading(title, level=3)
     add_para(f'鏃ユ湡锛歿date}', bold=True, size=9)
     add_para(desc, size=10)
     add_link_para('鏉ユ簮', url)
 
 # ============ 浜斻€佷細璁笌閫氱煡 ============
 doc.add_heading('浜斻€佷細璁笌閫氱煡', level=2)
 
 conferences = [
     ('ICAP 2026 (International Conference on Atomic Physics)',
      '2026骞?鏈堜笅鏃紙姝ｅ湪杩涜锛?, 'https://www.icap2026.org'),
     ('Quantum Sensing and Metrology (FIOLS 2026)',
      '2026骞?0鏈?, 'https://www.frontiersinoptics.com'),
     ('SPIE Photonics West 2027 鈥?Quantum Sensing',
      '鎽樿鎴锛?026骞?鏈?, 'https://spie.org'),
 ]
 
 for name, date, url in conferences:
     p = doc.add_paragraph()
     run = p.add_run(f'鈥?{name}')
     run.bold = True
     run.font.size = Pt(10)
     p.add_run(f' 鈥?{date}')
     p.paragraph_format.space_after = Pt(2)
     add_link_para('  浼氳缃戝潃', url, size=9)
 
 # ============ 鍏€佷笅鍛ㄩ鍒?============
 doc.add_heading('鍏€佷笅鍛ㄩ鍒?, level=2)
 add_para(
     '涓嬪懆灏嗙户缁窡韪噺瀛愬共娑変华锛堢壒鍒槸 SU(1,1) 鏋勫瀷鐨勫疄楠岃繘灞曪級銆佺儹鍘熷瓙閲忓瓙鍏夋簮涓庡瓨鍌ㄦ柟鍚戠殑'
     '鏈€鏂版枃鐚€傞噸鐐瑰叧娉ㄤ互涓嬩富棰橈細',
     size=10
 )
 outlooks = [
     '鐑師瀛愯捀姘斾腑鐨勫洓娉㈡贩棰戝帇缂╁厜婧愭渶鏂板疄楠岀粨鏋?,
     'OPM 纾佸姏璁″湪澶氶€氶亾鑴戠鍥?(MEG) 涓殑绯荤粺闆嗘垚杩涘睍',
     '閲忓瓙闄€铻轰华鐨勫櫔澹板垎鏋愬拰闀挎湡绋冲畾鎬у姣?,
     '缇庡浗 Farseer 璁″垝鍜?NSF Triad 鐨勫悗缁祫鍔╁叕鍛婂姩鎬?,
     'ICAP 2026 浼氳涓噺瀛愮簿瀵嗘祴閲忕浉鍏虫姤鍛?,
 ]
 for o in outlooks:
     p = doc.add_paragraph()
     p.style = doc.styles['List Bullet']
     run = p.add_run(o)
     run.font.size = Pt(10)
     p.paragraph_format.space_after = Pt(2)
 
 # ============ FOOTER ============
 doc.add_paragraph()
 footer_p = doc.add_paragraph()
 footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
 frun = footer_p.add_run('鈥斺€?鑷姩鍖栫敓鎴愪簬閲忓瓙鍏夊绮惧瘑娴嬮噺 路 鏂囩尞鏃ユ姤绯荤粺 鈥斺€?)
 frun.font.size = Pt(8)
 frun.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
 frun.italic = True
 
 # ============ SAVE ============
 out_dir = r'D:\everyday_recommand\reports\2026\07'
 os.makedirs(out_dir, exist_ok=True)
 out_path = os.path.join(out_dir, '2026-07-24-鍛ㄦ姤.docx')
 doc.save(out_path)
 print(f'DOCX saved: {out_path}')
 print(f'File size: {os.path.getsize(out_path)} bytes')
